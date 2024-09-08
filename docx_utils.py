from docx import Document
from io import BytesIO
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_rtl_paragraph(paragraph):
    """Set paragraph alignment and direction to Right-to-Left (RTL)."""
    paragraph.alignment = 2  # 2 corresponds to right-aligned in python-docx
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    rPr = run._r.get_or_add_rPr()
    rtl = OxmlElement('w:rtl')
    rtl.set(qn('w:val'), '1')
    rPr.append(rtl)

def create_docx_from_text(text_by_page, language):
    doc = Document()
    for page in text_by_page:
        para = doc.add_paragraph(page['text'] if page['text'] else 'No text found on this page.')
        if language == 'ara':
            set_rtl_paragraph(para)
        doc.add_page_break()
    return doc

def convert_docx_to_bytes(doc):
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()